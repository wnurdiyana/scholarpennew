"""ManuscriptForge — Q1-grade AI manuscript writer backend.

Endpoints (all prefixed with /api):
- Auth (JWT email/password): /auth/register, /auth/login, /auth/me, /auth/logout
- Auth (Emergent Google OAuth): /auth/google/session
- Manuscripts CRUD: /manuscripts, /manuscripts/{id}
- Section generation: /manuscripts/{id}/sections/{section_key}/generate
- References (Crossref): /references/search
- Export: /manuscripts/{id}/export
"""
from __future__ import annotations

import io
import logging
import os
import re
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import bcrypt
import httpx
import jwt
from dotenv import load_dotenv
from fastapi import APIRouter, Cookie, Depends, FastAPI, Header, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, EmailStr, Field
from starlette.middleware.cors import CORSMiddleware

# Optional integrations
from emergentintegrations.llm.chat import LlmChat, UserMessage  # type: ignore

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

# ---------- Config ----------
MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
EMERGENT_LLM_KEY = os.environ["EMERGENT_LLM_KEY"]
JWT_SECRET = os.environ["JWT_SECRET"]
JWT_ALG = "HS256"
JWT_EXPIRY_DAYS = 7

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="ManuscriptForge API")
api = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s :: %(message)s")
logger = logging.getLogger("manuscriptforge")


# ---------- Section catalog ----------
SECTION_KEYS = [
    "title",
    "abstract",
    "keywords",
    "introduction",
    "literature_review",
    "research_gap",
    "novelty",
    "methodology",
    "experimental_setup",
    "math_modeling",
    "ml_framework",
    "results",
    "discussion",
    "comparison",
    "practical_implications",
    "limitations",
    "conclusion",
    "future_work",
    "nomenclature",
    "references",
]

SECTION_LABELS = {
    "title": "Title (with 5 alternatives)",
    "abstract": "Abstract",
    "keywords": "Keywords",
    "introduction": "Introduction",
    "literature_review": "Literature Review",
    "research_gap": "Research Gap",
    "novelty": "Novelty Statement",
    "methodology": "Methodology",
    "experimental_setup": "Experimental Setup",
    "math_modeling": "Mathematical / Statistical Modeling",
    "ml_framework": "Machine Learning Framework",
    "results": "Results",
    "discussion": "Discussion",
    "comparison": "Comparison with Existing Literature",
    "practical_implications": "Practical Implications",
    "limitations": "Limitations",
    "conclusion": "Conclusion",
    "future_work": "Future Recommendations",
    "nomenclature": "Nomenclature",
    "references": "References",
}


# ---------- Models ----------
class RegisterIn(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)
    name: str = Field(min_length=1)


class LoginIn(BaseModel):
    email: EmailStr
    password: str


class GoogleSessionIn(BaseModel):
    session_id: str


class UserOut(BaseModel):
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    auth_provider: str


class ManuscriptInputs(BaseModel):
    title: Optional[str] = ""
    field: Optional[str] = ""
    abstract_idea: Optional[str] = ""
    objectives: Optional[str] = ""
    hypothesis: Optional[str] = ""
    methodology: Optional[str] = ""
    experimental_setup: Optional[str] = ""
    data_summary: Optional[str] = ""
    figures_tables: Optional[str] = ""
    statistical_results: Optional[str] = ""
    ml_models: Optional[str] = ""
    equations: Optional[str] = ""
    literature_sources: Optional[str] = ""
    journal_target: Optional[str] = ""
    keywords: Optional[str] = ""
    findings: Optional[str] = ""
    discussion_points: Optional[str] = ""
    limitations: Optional[str] = ""
    future_work: Optional[str] = ""
    citation_style: Optional[str] = "APA"


class ManuscriptCreateIn(BaseModel):
    inputs: ManuscriptInputs


class ManuscriptPatchIn(BaseModel):
    inputs: Optional[ManuscriptInputs] = None
    section_overrides: Optional[Dict[str, str]] = None


class SectionGenerateIn(BaseModel):
    extra_instructions: Optional[str] = ""


# ---------- Helpers ----------
def _now() -> datetime:
    return datetime.now(timezone.utc)


def _iso(dt: datetime) -> str:
    return dt.isoformat()


def _make_jwt(user_id: str) -> str:
    payload = {
        "sub": user_id,
        "iat": _now(),
        "exp": _now() + timedelta(days=JWT_EXPIRY_DAYS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


def _decode_jwt(token: str) -> Optional[str]:
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        return data.get("sub")
    except jwt.PyJWTError:
        return None


async def _user_from_session_token(token: str) -> Optional[Dict[str, Any]]:
    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not session:
        return None
    exp = session.get("expires_at")
    if isinstance(exp, str):
        exp = datetime.fromisoformat(exp)
    if exp and exp.tzinfo is None:
        exp = exp.replace(tzinfo=timezone.utc)
    if exp and exp < _now():
        return None
    return await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})


async def current_user(
    request: Request,
    authorization: Optional[str] = Header(default=None),
    session_token: Optional[str] = Cookie(default=None),
) -> Dict[str, Any]:
    """Resolve a user from either JWT (Bearer) or Emergent session_token (cookie or bearer)."""
    raw_token: Optional[str] = None
    if authorization and authorization.lower().startswith("bearer "):
        raw_token = authorization.split(" ", 1)[1].strip()
    if not raw_token and session_token:
        raw_token = session_token

    if not raw_token:
        raise HTTPException(status_code=401, detail="Not authenticated")

    # Try JWT first
    user_id = _decode_jwt(raw_token)
    if user_id:
        user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if user:
            return user

    # Fallback: Emergent session token
    user = await _user_from_session_token(raw_token)
    if user:
        return user

    raise HTTPException(status_code=401, detail="Invalid or expired session")


def _public_user(u: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "user_id": u["user_id"],
        "email": u.get("email", ""),
        "name": u.get("name", ""),
        "picture": u.get("picture"),
        "auth_provider": u.get("auth_provider", "password"),
    }


# ---------- Auth: JWT email/password ----------
@api.post("/auth/register")
async def register(payload: RegisterIn):
    existing = await db.users.find_one({"email": payload.email.lower()}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="An account with this email already exists")

    user_id = f"user_{uuid.uuid4().hex[:12]}"
    pw_hash = bcrypt.hashpw(payload.password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    doc = {
        "user_id": user_id,
        "email": payload.email.lower(),
        "name": payload.name.strip(),
        "picture": None,
        "password_hash": pw_hash,
        "auth_provider": "password",
        "created_at": _iso(_now()),
    }
    await db.users.insert_one(doc)
    token = _make_jwt(user_id)
    return {"token": token, "user": _public_user(doc)}


@api.post("/auth/login")
async def login(payload: LoginIn):
    user = await db.users.find_one({"email": payload.email.lower()}, {"_id": 0})
    if not user or not user.get("password_hash"):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    ok = bcrypt.checkpw(payload.password.encode("utf-8"), user["password_hash"].encode("utf-8"))
    if not ok:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = _make_jwt(user["user_id"])
    return {"token": token, "user": _public_user(user)}


@api.get("/auth/me")
async def me(user: Dict[str, Any] = Depends(current_user)):
    return _public_user(user)


@api.post("/auth/logout")
async def logout(authorization: Optional[str] = Header(default=None), session_token: Optional[str] = Cookie(default=None)):
    raw = None
    if authorization and authorization.lower().startswith("bearer "):
        raw = authorization.split(" ", 1)[1].strip()
    if not raw and session_token:
        raw = session_token
    if raw:
        await db.user_sessions.delete_one({"session_token": raw})
    return {"ok": True}


# ---------- Auth: Emergent Google OAuth ----------
@api.post("/auth/google/session")
async def google_session(payload: GoogleSessionIn):
    """Exchange Emergent session_id (from URL fragment) for an app session."""
    url = "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data"
    async with httpx.AsyncClient(timeout=20.0) as http:
        r = await http.get(url, headers={"X-Session-ID": payload.session_id})
    if r.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid Google session")
    data = r.json()
    email = (data.get("email") or "").lower()
    name = data.get("name") or email.split("@")[0]
    picture = data.get("picture")
    session_token = data.get("session_token")
    if not (email and session_token):
        raise HTTPException(status_code=502, detail="Bad upstream response")

    # Upsert user
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        user = {
            "user_id": user_id,
            "email": email,
            "name": name,
            "picture": picture,
            "auth_provider": "google",
            "password_hash": None,
            "created_at": _iso(_now()),
        }
        await db.users.insert_one(user)
    else:
        await db.users.update_one(
            {"user_id": user["user_id"]},
            {"$set": {"name": name, "picture": picture}},
        )
        user["name"] = name
        user["picture"] = picture

    expires = _now() + timedelta(days=7)
    await db.user_sessions.update_one(
        {"session_token": session_token},
        {"$set": {
            "user_id": user["user_id"],
            "session_token": session_token,
            "expires_at": _iso(expires),
            "created_at": _iso(_now()),
        }},
        upsert=True,
    )
    return {"token": session_token, "user": _public_user(user), "expires_at": _iso(expires)}


# ---------- Manuscripts ----------
def _empty_sections() -> Dict[str, Dict[str, Any]]:
    return {
        k: {"key": k, "label": SECTION_LABELS[k], "content": "", "status": "empty", "updated_at": None}
        for k in SECTION_KEYS
    }


@api.post("/manuscripts")
async def create_manuscript(payload: ManuscriptCreateIn, user: Dict[str, Any] = Depends(current_user)):
    mid = f"ms_{uuid.uuid4().hex[:14]}"
    doc = {
        "manuscript_id": mid,
        "user_id": user["user_id"],
        "title": payload.inputs.title or "Untitled Manuscript",
        "inputs": payload.inputs.model_dump(),
        "sections": _empty_sections(),
        "references": [],
        "created_at": _iso(_now()),
        "updated_at": _iso(_now()),
    }
    await db.manuscripts.insert_one(doc)
    return await db.manuscripts.find_one({"manuscript_id": mid}, {"_id": 0})


@api.get("/manuscripts")
async def list_manuscripts(user: Dict[str, Any] = Depends(current_user)):
    cursor = db.manuscripts.find({"user_id": user["user_id"]}, {"_id": 0, "sections": 0}).sort("updated_at", -1)
    items = await cursor.to_list(length=200)
    # Add a small summary count of generated sections
    out = []
    for it in items:
        full = await db.manuscripts.find_one({"manuscript_id": it["manuscript_id"]}, {"_id": 0, "sections": 1})
        generated = sum(1 for s in (full or {}).get("sections", {}).values() if s.get("status") == "complete")
        it["generated_sections"] = generated
        it["total_sections"] = len(SECTION_KEYS)
        out.append(it)
    return out


@api.get("/manuscripts/{mid}")
async def get_manuscript(mid: str, user: Dict[str, Any] = Depends(current_user)):
    doc = await db.manuscripts.find_one({"manuscript_id": mid, "user_id": user["user_id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Manuscript not found")
    return doc


@api.patch("/manuscripts/{mid}")
async def update_manuscript(mid: str, payload: ManuscriptPatchIn, user: Dict[str, Any] = Depends(current_user)):
    doc = await db.manuscripts.find_one({"manuscript_id": mid, "user_id": user["user_id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Manuscript not found")

    updates: Dict[str, Any] = {"updated_at": _iso(_now())}
    if payload.inputs is not None:
        updates["inputs"] = payload.inputs.model_dump()
        if payload.inputs.title:
            updates["title"] = payload.inputs.title

    if payload.section_overrides:
        sections = doc["sections"]
        for k, content in payload.section_overrides.items():
            if k in sections:
                sections[k]["content"] = content
                sections[k]["status"] = "complete" if content.strip() else "empty"
                sections[k]["updated_at"] = _iso(_now())
        updates["sections"] = sections

    await db.manuscripts.update_one({"manuscript_id": mid}, {"$set": updates})
    return await db.manuscripts.find_one({"manuscript_id": mid}, {"_id": 0})


@api.delete("/manuscripts/{mid}")
async def delete_manuscript(mid: str, user: Dict[str, Any] = Depends(current_user)):
    res = await db.manuscripts.delete_one({"manuscript_id": mid, "user_id": user["user_id"]})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Manuscript not found")
    return {"ok": True}


# ---------- Section generation (Claude Sonnet 4.5) ----------
SYSTEM_PROMPT = """You are an expert academic manuscript writer with deep expertise across engineering, maritime studies, environmental science, energy systems, machine learning, sustainability, and multidisciplinary scientific research. You write at the standard of top-tier Q1 WoS-indexed journals (Elsevier, Springer Nature, Taylor & Francis, Wiley, MDPI).

Your writing must be:
- Scientifically rigorous, logically structured, highly analytical, evidence-based.
- Formal academic English; varied sentence complexity; non-robotic, human-like prose.
- Free from generic filler, exaggerated claims, vague wording ("very good", "interesting", "amazing").
- Quantitative whenever possible; technically accurate; methodologically defensible.
- Resistant to AI-detection through natural scholarly cadence and reasoned argumentation.

Hard rules:
- Never fabricate DOIs, journal volumes, or article identifiers. If a specific reference is unverifiable, use a clearly marked placeholder like [Author, Year — verify DOI].
- Never invent unrealistic experimental numbers. Where the user did not provide quantitative findings, give clearly indicated illustrative ranges and label them as "indicative".
- Prefer 2021–2026 literature; foundational works permitted when cited as such.
- Do not use undefined abbreviations; expand on first use.
- Output clean Markdown only. Do not wrap output in code fences. Do not preface or append meta-commentary.
"""


def _format_inputs_block(inputs: Dict[str, Any]) -> str:
    lines = []
    for k, v in inputs.items():
        if v and str(v).strip():
            label = k.replace("_", " ").title()
            lines.append(f"- **{label}**: {v}")
    return "\n".join(lines) if lines else "_(No structured inputs provided — infer reasonably and clearly mark assumptions.)_"


def _section_prompt(section_key: str, manuscript: Dict[str, Any], extra: str) -> str:
    inputs = manuscript.get("inputs", {})
    title_hint = inputs.get("title") or manuscript.get("title") or "(untitled)"
    field = inputs.get("field") or "(unspecified field)"
    journal = inputs.get("journal_target") or "a high-impact Q1 WoS-indexed journal"
    citation_style = inputs.get("citation_style") or "APA"

    inputs_block = _format_inputs_block(inputs)

    # Existing sections for coherence
    sections = manuscript.get("sections", {})
    context_parts = []
    coherence_keys = ["title", "abstract", "introduction", "methodology", "results"]
    for ck in coherence_keys:
        if ck == section_key:
            continue
        existing = sections.get(ck, {}).get("content", "")
        if existing and existing.strip():
            context_parts.append(f"### Existing {SECTION_LABELS[ck]} (for coherence — do not repeat verbatim)\n{existing[:1500]}")
    context_block = "\n\n".join(context_parts) if context_parts else "_(No prior sections generated yet.)_"

    section_specs = {
        "title": """Generate ONE primary research title plus FIVE ranked alternatives.
Constraints: Q1-journal worthy; reveal methodology + objective; high-impact keywords; concise; no filler.
Format:
**Primary Title**
[your title]

**Alternative Titles (ranked)**
1. ...
2. ...
3. ...
4. ...
5. ...
""",
        "abstract": """Produce a 200–250 word structured Abstract covering: (1) background/problem, (2) gap, (3) objective, (4) methodology, (5) key quantitative findings, (6) scientific significance, (7) practical contribution. No citations. No undefined abbreviations. No bullet points — single flowing paragraph.""",
        "keywords": """Provide 5–8 high-impact keywords aligned with WoS/Scopus indexing. Comma-separated, no commentary.""",
        "introduction": """Write a critical, analytical Introduction (~800–1100 words). Establish global significance; cover recent (2021–2026) scientific developments; critically analyze prior work; expose unresolved problems; build a tight justification; explicitly state objectives and contributions in the final paragraph. Cite as [Author, Year] inline; verifiable references will be compiled in the References section.""",
        "literature_review": """Write a critical Literature Review (~900–1300 words). Synthesize, do not list. Group studies thematically. Compare strengths vs limitations. End with a sub-heading "### Comparative Synthesis" containing a Markdown comparison table with columns: Reference | Approach | Dataset / Setting | Key Finding | Limitation. Include at least 8 rows. Use inline [Author, Year] citations.""",
        "research_gap": """Identify research gaps in 250–400 words across explicit sub-headings: Technical, Methodological, Data, Industrial, Regional, and AI/Modeling (if applicable). Each gap must be specific and justify the present study.""",
        "novelty": """Write a Novelty Statement of 180–280 words. Make it specific, measurable, defensible, and publication-worthy. Use 4–6 numbered novelty points.""",
        "methodology": """Write a reproducible Methodology (~900–1200 words) covering: materials, equipment/software, datasets, preprocessing, calibration, algorithms, statistical techniques, uncertainty analysis, validation strategy. Use clear sub-headings (###). Include at least one equation rendered in LaTeX-style ($...$ or $$...$$).""",
        "experimental_setup": """Describe the Experimental Setup (~500–800 words). Include apparatus, instrumentation, operating conditions, sampling protocol, repeatability, and applicable standards. Add a Markdown table summarizing key parameters and ranges.""",
        "math_modeling": """Write a Mathematical / Statistical Modeling section (~600–900 words). Define variables; present governing equations in LaTeX (numbered (1), (2), ...); state assumptions; describe solver / estimation procedure; note convergence criteria.""",
        "ml_framework": """Write a Machine Learning Framework section (~600–900 words). Cover dataset preparation, train/validation/test split, architecture, hyperparameters (in a table), optimizer, activation functions, loss, regularization / overfitting prevention, and performance metrics (RMSE, MAE, R², F1 — as applicable). If ML is not applicable to the study, return a single short paragraph stating ML is not used and explaining why.""",
        "results": """Write the Results section (~900–1200 words). Present quantitative trends; reference figures/tables as Fig. 1, Fig. 2, Table 1 with caption suggestions in italics; include at least one results table (Markdown). Avoid pure description — explain mechanisms briefly; defer deeper interpretation to Discussion.""",
        "discussion": """Write a deep, mechanistic Discussion (~900–1300 words). Explain *why* observed trends occur; connect to theory; benchmark against prior literature with [Author, Year] citations; discuss engineering, sustainability, and policy implications. Avoid restating results.""",
        "comparison": """Write a Comparison with Existing Literature section (~400–600 words) plus a Markdown table: Study | Method | Reported Metric | This Study | Δ Improvement. Cite [Author, Year] inline. End with 2–3 sentences on why the present results advance the state of the art.""",
        "practical_implications": """Write Practical Implications (~300–500 words). Address industry adoption, regulatory relevance, operational deployment, scalability, and economic considerations.""",
        "limitations": """Write Limitations (~250–400 words). Be candid: data scope, model assumptions, generalizability, computational cost, external validity.""",
        "conclusion": """Write a Conclusion (400–700 words). Summarize major quantitative findings; emphasize the contribution and significance; outline implications; acknowledge limitations briefly; close with future-work hooks. No new citations.""",
        "future_work": """Write Future Recommendations (~250–400 words) as numbered, specific, feasible research directions tied to gaps identified in Limitations.""",
        "nomenclature": """Produce a Nomenclature section as a Markdown table with columns: Symbol | Description | Unit. Group by Roman, Greek, Subscripts, Superscripts, Abbreviations using ### sub-headings. Include only symbols actually used in this manuscript's other sections.""",
        "references": f"""Compile a References section in **{citation_style}** style (numbered or alphabetized as per style). Prioritize 2021–2026 Q1/Q2 WoS- or Scopus-indexed sources; ~10–20% foundational. Include DOI where it can be reasonably stated. Where a DOI cannot be verified, append " — verify DOI" after the entry. Provide 25–45 entries. Do NOT fabricate DOIs; mark unverifiable ones explicitly. Do not include commentary; output only the reference list.""",
    }

    spec = section_specs.get(section_key, "Write this section to Q1 journal standard.")

    return f"""# Task
Generate the **{SECTION_LABELS[section_key]}** section of a manuscript targeted at **{journal}** in the field of **{field}**.
Working title: *{title_hint}*. Citation style: **{citation_style}**.

# Author-supplied research inputs
{inputs_block}

# Coherence context (excerpts of already-generated sections)
{context_block}

# Section-specific instructions
{spec}

# Additional instructions from author
{extra or '_(none)_'}

# Output
Return clean Markdown for THIS section only. Do NOT include the section heading itself; start directly with the body content."""


@api.post("/manuscripts/{mid}/sections/{section_key}/generate")
async def generate_section(
    mid: str,
    section_key: str,
    payload: SectionGenerateIn,
    user: Dict[str, Any] = Depends(current_user),
):
    if section_key not in SECTION_KEYS:
        raise HTTPException(status_code=400, detail="Unknown section key")
    doc = await db.manuscripts.find_one({"manuscript_id": mid, "user_id": user["user_id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Manuscript not found")

    # Mark generating
    await db.manuscripts.update_one(
        {"manuscript_id": mid},
        {"$set": {
            f"sections.{section_key}.status": "generating",
            f"sections.{section_key}.updated_at": _iso(_now()),
            "updated_at": _iso(_now()),
        }},
    )

    prompt = _section_prompt(section_key, doc, payload.extra_instructions or "")

    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"{mid}:{section_key}:{uuid.uuid4().hex[:6]}",
            system_message=SYSTEM_PROMPT,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        response_text = await chat.send_message(UserMessage(text=prompt))
        if not isinstance(response_text, str):
            response_text = str(response_text)
        response_text = response_text.strip()
        # Strip accidental code fences
        if response_text.startswith("```"):
            response_text = re.sub(r"^```[a-zA-Z]*\n?", "", response_text)
            response_text = re.sub(r"\n?```$", "", response_text)
    except Exception as exc:
        logger.exception("LLM generation failed")
        await db.manuscripts.update_one(
            {"manuscript_id": mid},
            {"$set": {f"sections.{section_key}.status": "error"}},
        )
        raise HTTPException(status_code=502, detail=f"Generation failed: {exc}")

    await db.manuscripts.update_one(
        {"manuscript_id": mid},
        {"$set": {
            f"sections.{section_key}.content": response_text,
            f"sections.{section_key}.status": "complete",
            f"sections.{section_key}.updated_at": _iso(_now()),
            "updated_at": _iso(_now()),
        }},
    )

    # If title section, also update top-level title for dashboard listing
    if section_key == "title":
        m = re.search(r"\*\*Primary Title\*\*\s*\n+([^\n]+)", response_text)
        if m:
            new_title = m.group(1).strip().lstrip("*# ").strip()
            if new_title:
                await db.manuscripts.update_one(
                    {"manuscript_id": mid},
                    {"$set": {"title": new_title[:200]}},
                )

    fresh = await db.manuscripts.find_one({"manuscript_id": mid}, {"_id": 0})
    return {"section": fresh["sections"][section_key], "manuscript_title": fresh.get("title")}


# ---------- References (Crossref) ----------
@api.get("/references/search")
async def reference_search(q: str = Query(..., min_length=2), rows: int = 10, user: Dict[str, Any] = Depends(current_user)):
    url = "https://api.crossref.org/works"
    params = {
        "query": q,
        "rows": min(max(rows, 1), 25),
        "select": "DOI,title,author,issued,container-title,volume,issue,page,publisher,type,abstract,URL",
        "sort": "relevance",
    }
    try:
        async with httpx.AsyncClient(timeout=15.0) as http:
            r = await http.get(url, params=params, headers={"User-Agent": "ManuscriptForge/1.0 (mailto:research@manuscriptforge.app)"})
        r.raise_for_status()
        items = r.json().get("message", {}).get("items", [])
    except Exception as exc:
        logger.warning("Crossref error: %s", exc)
        raise HTTPException(status_code=502, detail="Reference search unavailable")

    out = []
    for it in items:
        authors = []
        for a in it.get("author", [])[:6]:
            given = a.get("given", "")
            family = a.get("family", "")
            full = (f"{given} {family}").strip()
            if full:
                authors.append(full)
        year = None
        issued = it.get("issued", {}).get("date-parts", [[None]])
        if issued and issued[0]:
            year = issued[0][0]
        out.append({
            "doi": it.get("DOI"),
            "title": (it.get("title") or [""])[0],
            "authors": authors,
            "year": year,
            "container": (it.get("container-title") or [""])[0],
            "volume": it.get("volume"),
            "issue": it.get("issue"),
            "page": it.get("page"),
            "publisher": it.get("publisher"),
            "type": it.get("type"),
            "url": it.get("URL"),
        })
    return {"results": out}


# ---------- Export ----------
def _assemble_markdown(doc: Dict[str, Any]) -> str:
    title = doc.get("title") or "Untitled Manuscript"
    parts: List[str] = [f"# {title}\n"]
    sections = doc.get("sections", {})
    for k in SECTION_KEYS:
        s = sections.get(k, {})
        content = (s.get("content") or "").strip()
        if not content:
            continue
        parts.append(f"\n\n## {SECTION_LABELS[k]}\n\n{content}\n")
    return "".join(parts)


def _build_docx(doc: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    document.add_heading(doc.get("title") or "Untitled Manuscript", level=0)

    sections = doc.get("sections", {})
    for k in SECTION_KEYS:
        s = sections.get(k, {})
        content = (s.get("content") or "").strip()
        if not content:
            continue
        document.add_heading(SECTION_LABELS[k], level=1)
        # Simple paragraph splitting; preserve tables as raw markdown text blocks.
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                document.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                document.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                document.add_heading(block[2:].strip(), level=1)
            else:
                document.add_paragraph(block)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_pdf(doc: Dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName="Times-Bold", fontSize=18, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName="Times-Bold", fontSize=14, spaceAfter=8)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="Times-Roman", fontSize=11, leading=15, spaceAfter=8)

    story = [Paragraph(doc.get("title") or "Untitled Manuscript", h1), Spacer(1, 0.4 * cm)]
    sections = doc.get("sections", {})
    for k in SECTION_KEYS:
        s = sections.get(k, {})
        content = (s.get("content") or "").strip()
        if not content:
            continue
        story.append(Paragraph(SECTION_LABELS[k], h2))
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # Escape minimal HTML chars
            safe = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Convert markdown bold
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
            safe = re.sub(r"\*(.+?)\*", r"<i>\1</i>", safe)
            safe = safe.replace("\n", "<br/>")
            story.append(Paragraph(safe, body))

    pdf.build(story)
    return buf.getvalue()


@api.get("/manuscripts/{mid}/export")
async def export_manuscript(mid: str, format: str = Query("md"), user: Dict[str, Any] = Depends(current_user)):
    doc = await db.manuscripts.find_one({"manuscript_id": mid, "user_id": user["user_id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Manuscript not found")

    fmt = format.lower()
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", (doc.get("title") or "manuscript"))[:60] or "manuscript"

    if fmt in ("md", "markdown"):
        content = _assemble_markdown(doc).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.md"'},
        )
    if fmt == "docx":
        data = _build_docx(doc)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    if fmt == "pdf":
        data = _build_pdf(doc)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )
    raise HTTPException(status_code=400, detail="Unsupported format. Use md, docx, or pdf.")


# ---------- Misc ----------
@api.get("/")
async def root():
    return {"app": "ManuscriptForge", "ok": True, "sections": SECTION_KEYS}


@api.get("/sections/catalog")
async def sections_catalog():
    return [{"key": k, "label": SECTION_LABELS[k]} for k in SECTION_KEYS]


app.include_router(api)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
