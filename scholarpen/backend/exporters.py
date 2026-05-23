"""Manuscript exporters with hierarchical numbering and proper subheading/table rendering.

Sections are numbered 1, 2, 3, … in the order they appear.
Inside each section, Markdown sub-headings are numbered hierarchically:
  ## Heading   ->  1.1, 1.2, …
  ### Heading  ->  1.1.1, 1.1.2, …
  #### Heading ->  1.1.1.1, …

Supports:
- paragraphs (with **bold**, *italic*, inline `code`)
- bullet lists (-, *) and ordered lists (1.)
- Markdown tables (| col | col |)
- inline citations / brackets preserved as text
"""
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Tuple

# ---------- Markdown -> block parser ----------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_DIV_RE = re.compile(r"^\s*\|?\s*:?-{2,}.*$")


def parse_markdown_blocks(text: str) -> List[Dict[str, Any]]:
    """Convert markdown text to a flat list of structured blocks."""
    if not text:
        return []
    lines = text.replace("\r\n", "\n").split("\n")
    blocks: List[Dict[str, Any]] = []
    i = 0
    n = len(lines)

    def flush_paragraph(buf: List[str]) -> None:
        joined = " ".join(s.strip() for s in buf if s.strip())
        if joined:
            blocks.append({"type": "paragraph", "text": joined})

    para_buf: List[str] = []

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Blank line -> paragraph break
        if not stripped:
            flush_paragraph(para_buf)
            para_buf = []
            i += 1
            continue

        # Heading
        m = _HEADING_RE.match(stripped)
        if m:
            flush_paragraph(para_buf)
            para_buf = []
            level = len(m.group(1))
            blocks.append({"type": "heading", "level": level, "text": m.group(2).strip()})
            i += 1
            continue

        # Table: current line and next line look like a table header + divider
        if _TABLE_ROW_RE.match(line) and (i + 1 < n) and _TABLE_DIV_RE.match(lines[i + 1]):
            flush_paragraph(para_buf)
            para_buf = []
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows: List[List[str]] = []
            j = i + 2
            while j < n and _TABLE_ROW_RE.match(lines[j]):
                row = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                # pad/trim to header length
                if len(row) < len(header):
                    row += [""] * (len(header) - len(row))
                elif len(row) > len(header):
                    row = row[: len(header)]
                rows.append(row)
                j += 1
            blocks.append({"type": "table", "headers": header, "rows": rows})
            i = j
            continue

        # Lists: collect consecutive list items of same kind, allowing blank lines between items.
        if _BULLET_RE.match(stripped) or _ORDERED_RE.match(stripped):
            flush_paragraph(para_buf)
            para_buf = []
            ordered = bool(_ORDERED_RE.match(stripped))
            items: List[str] = []
            while i < n:
                s = lines[i].strip()
                if not s:
                    # Peek ahead: if the next non-blank line is also a same-kind list item, keep going.
                    j = i + 1
                    while j < n and not lines[j].strip():
                        j += 1
                    if j < n:
                        next_s = lines[j].strip()
                        same_kind = (_ORDERED_RE.match(next_s) if ordered else _BULLET_RE.match(next_s))
                        if same_kind:
                            i = j
                            continue
                    break
                mm = _ORDERED_RE.match(s) if ordered else _BULLET_RE.match(s)
                if not mm:
                    # mixed kinds → break list
                    if (ordered and _BULLET_RE.match(s)) or (not ordered and _ORDERED_RE.match(s)):
                        break
                    if _HEADING_RE.match(s) or _TABLE_ROW_RE.match(lines[i]):
                        break
                    # continuation line -> append to last item
                    if items:
                        items[-1] = items[-1] + " " + s
                        i += 1
                        continue
                    break
                items.append(mm.group(1).strip())
                i += 1
            blocks.append({"type": "list", "ordered": ordered, "items": items})
            continue

        # Default: paragraph accumulator
        para_buf.append(stripped)
        i += 1

    flush_paragraph(para_buf)
    return blocks


# ---------- Numbering ----------

def compute_numbering(top_index: int, sub_levels: List[int]) -> str:
    """Produce dotted numbering string e.g. '1', '1.1', '1.2.3'."""
    parts = [str(top_index)] + [str(x) for x in sub_levels]
    return ".".join(parts)


class HierarchyCounter:
    """Tracks 1.x.x.x… numbering as we iterate over a section's headings.

    Strategy: keep a sparse map {level: count}. On a new heading at level L,
    drop any counters at levels deeper than L and increment counters[L].
    Numbering is `top . count[L1] . count[L2] . … . count[L]` for every
    level <= L that has been observed. This collapses correctly when a
    section's first sub-heading is `###` (no parent `##`).
    """

    def __init__(self, top_index: int) -> None:
        self.top = top_index
        self.counters: Dict[int, int] = {}

    def assign(self, level: int) -> str:
        # Drop any deeper levels we previously tracked.
        for lvl in list(self.counters.keys()):
            if lvl > level:
                del self.counters[lvl]
        self.counters[level] = self.counters.get(level, 0) + 1
        parts: List[str] = [str(self.top)]
        for lvl in sorted(self.counters.keys()):
            if lvl > level:
                break
            parts.append(str(self.counters[lvl]))
        return ".".join(parts)


# ---------- Inline markdown -> reportlab/docx friendly ----------

_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)([^\*\n]+?)\*(?!\*)")
_INLINE_CODE = re.compile(r"`([^`\n]+?)`")


def _inline_to_reportlab(text: str) -> str:
    """Convert inline markdown to reportlab Paragraph HTML."""
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    safe = _INLINE_BOLD.sub(r"<b>\1</b>", safe)
    safe = _INLINE_ITALIC.sub(r"<i>\1</i>", safe)
    safe = _INLINE_CODE.sub(r'<font face="Courier">\1</font>', safe)
    return safe


def _inline_runs_for_docx(text: str) -> List[Tuple[str, Dict[str, bool]]]:
    """Return a list of (text, {bold, italic, code}) tuples for docx runs."""
    runs: List[Tuple[str, Dict[str, bool]]] = []
    # Process bold first, then italic — using simple non-nested handling sufficient for our content.
    # Tokenize bold
    pos = 0
    tokens: List[Tuple[str, Dict[str, bool]]] = []
    for m in _INLINE_BOLD.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], {}))
        tokens.append((m.group(1), {"bold": True}))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], {}))

    # Now apply italic to each non-bold token
    for chunk, fmt in tokens:
        if fmt.get("bold"):
            runs.append((chunk, fmt))
            continue
        ipos = 0
        for im in _INLINE_ITALIC.finditer(chunk):
            if im.start() > ipos:
                runs.append((chunk[ipos:im.start()], {}))
            runs.append((im.group(1), {"italic": True}))
            ipos = im.end()
        if ipos < len(chunk):
            runs.append((chunk[ipos:], {}))
    return runs


# ---------- Section iteration with numbering ----------

def iter_numbered_sections(doc: Dict[str, Any], section_keys: List[str], section_labels: Dict[str, str]):
    """Yield (top_index, label, content_blocks, hierarchy_counter) for each non-empty section."""
    sections = doc.get("sections", {})
    top = 0
    for key in section_keys:
        s = sections.get(key, {})
        content = (s.get("content") or "").strip()
        if not content:
            continue
        top += 1
        blocks = parse_markdown_blocks(content)
        yield top, section_labels[key], blocks, HierarchyCounter(top)


# ---------- Markdown export with numbering ----------

def assemble_markdown(doc: Dict[str, Any], section_keys: List[str], section_labels: Dict[str, str]) -> str:
    title = doc.get("title") or "Untitled Manuscript"
    out: List[str] = [f"# {title}", ""]
    for top, label, blocks, counter in iter_numbered_sections(doc, section_keys, section_labels):
        out.append("")
        out.append(f"## {top}. {label}")
        out.append("")
        # Reset counter for this section, then re-emit each block with numbering.
        for b in blocks:
            if b["type"] == "heading":
                level = b["level"]
                if level <= 1:
                    # treat lone # inside a section as a paragraph to avoid clashing with section numbering
                    out.append(b["text"])
                    out.append("")
                    continue
                num = counter.assign(level)
                hashes = "#" * min(6, level + 1)  # shift down 1 so section is ##
                out.append(f"{hashes} {num} {b['text']}")
                out.append("")
            elif b["type"] == "paragraph":
                out.append(b["text"])
                out.append("")
            elif b["type"] == "list":
                for j, item in enumerate(b["items"], start=1):
                    prefix = f"{j}." if b["ordered"] else "-"
                    out.append(f"{prefix} {item}")
                out.append("")
            elif b["type"] == "table":
                out.append("| " + " | ".join(b["headers"]) + " |")
                out.append("| " + " | ".join(["---"] * len(b["headers"])) + " |")
                for row in b["rows"]:
                    out.append("| " + " | ".join(row) + " |")
                out.append("")
    return "\n".join(out).rstrip() + "\n"


# ---------- DOCX export ----------

def build_docx(doc: Dict[str, Any], section_keys: List[str], section_labels: Dict[str, str]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    # Manuscript title
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(doc.get("title") or "Untitled Manuscript")
    run.bold = True
    run.font.size = Pt(20)
    document.add_paragraph()

    def add_inline_paragraph(text: str, *, italic: bool = False) -> None:
        p = document.add_paragraph()
        runs = _inline_runs_for_docx(text)
        for chunk, fmt in runs:
            r = p.add_run(chunk)
            r.bold = fmt.get("bold", False)
            r.italic = fmt.get("italic", False) or italic

    def add_heading_text(level: int, numbering: str, text: str) -> None:
        """Add a numbered heading at the given Word heading level."""
        para = document.add_paragraph()
        para.style = document.styles[f"Heading {min(level, 9)}"]
        r = para.add_run(f"{numbering}  {text}" if numbering else text)
        r.bold = True

    for top, label, blocks, counter in iter_numbered_sections(doc, section_keys, section_labels):
        add_heading_text(1, f"{top}.", label)

        for b in blocks:
            if b["type"] == "heading":
                level = b["level"]
                if level <= 1:
                    add_inline_paragraph(b["text"])
                    continue
                num = counter.assign(level)
                add_heading_text(min(level + 1, 4), num, b["text"])
            elif b["type"] == "paragraph":
                add_inline_paragraph(b["text"])
            elif b["type"] == "list":
                for j, item in enumerate(b["items"], start=1):
                    style_name = "List Number" if b["ordered"] else "List Bullet"
                    try:
                        p = document.add_paragraph(style=style_name)
                    except KeyError:
                        p = document.add_paragraph()
                        p.add_run(f"{j}. " if b["ordered"] else "• ").bold = False
                    for chunk, fmt in _inline_runs_for_docx(item):
                        r = p.add_run(chunk)
                        r.bold = fmt.get("bold", False)
                        r.italic = fmt.get("italic", False)
            elif b["type"] == "table":
                headers = b["headers"]
                rows = b["rows"]
                if not headers:
                    continue
                table = document.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                hdr_cells = table.rows[0].cells
                for ci, h in enumerate(headers):
                    cell = hdr_cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(h)
                    run.bold = True
                for ri, row in enumerate(rows, start=1):
                    cells = table.rows[ri].cells
                    for ci, val in enumerate(row):
                        cells[ci].text = ""
                        p = cells[ci].paragraphs[0]
                        for chunk, fmt in _inline_runs_for_docx(val):
                            run = p.add_run(chunk)
                            run.bold = fmt.get("bold", False)
                            run.italic = fmt.get("italic", False)
                document.add_paragraph()  # spacer

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------- PDF export ----------

def build_pdf(doc: Dict[str, Any], section_keys: List[str], section_labels: Dict[str, str]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
        title=doc.get("title") or "Manuscript",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "manuscript_title",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=20,
        leading=24,
        alignment=1,  # center
        spaceAfter=18,
    )
    h_styles = {
        1: ParagraphStyle("h1", parent=styles["Heading1"], fontName="Times-Bold", fontSize=15, leading=19, spaceBefore=14, spaceAfter=8),
        2: ParagraphStyle("h2", parent=styles["Heading2"], fontName="Times-Bold", fontSize=12.5, leading=16, spaceBefore=10, spaceAfter=6),
        3: ParagraphStyle("h3", parent=styles["Heading3"], fontName="Times-BoldItalic", fontSize=11.5, leading=15, spaceBefore=8, spaceAfter=4),
        4: ParagraphStyle("h4", parent=styles["Heading4"], fontName="Times-Italic", fontSize=11, leading=14, spaceBefore=6, spaceAfter=4),
    }
    body_style = ParagraphStyle(
        "body", parent=styles["BodyText"], fontName="Times-Roman", fontSize=11, leading=15, spaceAfter=6, alignment=4  # justify
    )
    list_item_style = ParagraphStyle("li", parent=body_style, spaceAfter=2)

    story = []
    story.append(Paragraph(_inline_to_reportlab(doc.get("title") or "Untitled Manuscript"), title_style))

    for top, label, blocks, counter in iter_numbered_sections(doc, section_keys, section_labels):
        story.append(Paragraph(f"<b>{top}.&nbsp;&nbsp;{_inline_to_reportlab(label)}</b>", h_styles[1]))

        for b in blocks:
            if b["type"] == "heading":
                level = b["level"]
                if level <= 1:
                    story.append(Paragraph(_inline_to_reportlab(b["text"]), body_style))
                    continue
                num = counter.assign(level)
                style = h_styles.get(min(level, 4), h_styles[4])
                story.append(Paragraph(f"<b>{num}&nbsp;&nbsp;{_inline_to_reportlab(b['text'])}</b>", style))
            elif b["type"] == "paragraph":
                story.append(Paragraph(_inline_to_reportlab(b["text"]), body_style))
            elif b["type"] == "list":
                items = [ListItem(Paragraph(_inline_to_reportlab(it), list_item_style), leftIndent=14) for it in b["items"]]
                story.append(ListFlowable(
                    items,
                    bulletType="1" if b["ordered"] else "bullet",
                    start="1" if b["ordered"] else None,
                    leftIndent=18,
                ))
                story.append(Spacer(1, 4))
            elif b["type"] == "table":
                headers = b["headers"]
                rows = b["rows"]
                if not headers:
                    continue
                data = [[Paragraph(f"<b>{_inline_to_reportlab(h)}</b>", body_style) for h in headers]]
                for row in rows:
                    data.append([Paragraph(_inline_to_reportlab(c), body_style) for c in row])
                tbl = Table(data, repeatRows=1, hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f4f4f5")),
                    ("LINEBELOW", (0, 0), (-1, 0), 1.2, colors.HexColor("#18181b")),
                    ("LINEBELOW", (0, 1), (-1, -1), 0.3, colors.HexColor("#e4e4e7")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 8))

    pdf.build(story)
    return buf.getvalue()
